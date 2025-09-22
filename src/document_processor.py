"""
Document processing module for CLM automation system.
Handles document ingestion, content extraction, and preprocessing.
"""

import os
import logging
import re
from typing import Dict, Any, List, Optional
from datetime import datetime
import json

# File processing imports
import PyPDF2
from docx import Document
import pytesseract
from pdf2image import convert_from_path
from PIL import Image

# LangChain imports
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangchainDocument

from src.config import Config
from src.database import DatabaseManager
from src.embeddings import EmbeddingManager

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Processes documents for ingestion into the CLM system"""
    
    def __init__(self, db_manager: DatabaseManager):
        self.db_manager = db_manager
        self.embedding_manager = EmbeddingManager(db_manager)
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=Config.CHUNK_SIZE,
            chunk_overlap=Config.CHUNK_OVERLAP,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
    
    def process_folder(self, folder_path: str = Config.DOCUMENTS_FOLDER) -> Dict[str, Any]:
        """Process all documents in a folder"""
        logger.info(f"Processing documents in folder: {folder_path}")
        
        results = {
            "processed": [],
            "failed": [],
            "total_documents": 0,
            "total_chunks": 0
        }
        
        if not os.path.exists(folder_path):
            logger.error(f"Folder not found: {folder_path}")
            return results
        
        # Get all supported files
        supported_files = self._get_supported_files(folder_path)
        results["total_documents"] = len(supported_files)
        
        for file_path in supported_files:
            try:
                result = self.process_document(file_path)
                if result:
                    results["processed"].append({
                        "filename": os.path.basename(file_path),
                        "document_id": result["document_id"],
                        "chunks": result["chunk_count"]
                    })
                    results["total_chunks"] += result["chunk_count"]
                else:
                    results["failed"].append(os.path.basename(file_path))
            except Exception as e:
                logger.error(f"Failed to process {file_path}: {e}")
                results["failed"].append(os.path.basename(file_path))
        
        logger.info(f"Processing complete. Processed: {len(results['processed'])}, "
                   f"Failed: {len(results['failed'])}, Total chunks: {results['total_chunks']}")
        
        return results
    
    def process_document(self, file_path: str) -> Optional[Dict[str, Any]]:
        """Process a single document"""
        try:
            logger.info(f"Processing document: {file_path}")
            
            filename = os.path.basename(file_path)
            file_extension = os.path.splitext(filename)[1].lower()
            
            # Extract content based on file type
            content, metadata = self._extract_content(file_path, file_extension)
            
            if not content.strip():
                logger.warning(f"No content extracted from {filename}")
                return None
            
            # Store document in database
            document_id = self.db_manager.insert_document(
                filename=filename,
                file_type=file_extension[1:],  # Remove the dot
                content=content,
                metadata=metadata
            )
            
            # Process chunks and create embeddings
            chunks = self._create_chunks(content, filename)
            chunk_count = len(chunks)
            
            # Generate and store embeddings
            if chunks:
                embedding_success = self.embedding_manager.generate_and_store_embeddings(document_id, chunks)
                if not embedding_success:
                    logger.warning(f"Failed to generate embeddings for {filename}")
            
            # Extract and store contract information
            contract_data = self._extract_contract_info(content, filename)
            if contract_data:
                self.db_manager.insert_contract(document_id, contract_data)
            
            logger.info(f"Successfully processed {filename}: {chunk_count} chunks created")
            
            return {
                "document_id": document_id,
                "filename": filename,
                "chunk_count": chunk_count,
                "contract_data": contract_data
            }
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            return None
    
    def process_single_document(self, file_path: str, filename: str = None, 
                              extract_contracts: bool = True, 
                              custom_chunk_size: int = None) -> Dict[str, Any]:
        """Process a single uploaded document with custom options
        
        Args:
            file_path: Path to the document file
            filename: Custom filename (if different from file_path)
            extract_contracts: Whether to extract contract information
            custom_chunk_size: Custom chunk size for text splitting
            
        Returns:
            Dict containing processing results
        """
        try:
            if filename is None:
                filename = os.path.basename(file_path)
            
            logger.info(f"Processing uploaded document: {filename}")
            
            file_extension = os.path.splitext(filename)[1].lower()
            
            # Extract content based on file type
            content, metadata = self._extract_content(file_path, file_extension)
            
            if not content.strip():
                logger.warning(f"No content extracted from {filename}")
                return {
                    "success": False,
                    "error": "No content could be extracted from the document",
                    "filename": filename
                }
            
            # Store document in database
            document_id = self.db_manager.insert_document(
                filename=filename,
                file_type=file_extension[1:],  # Remove the dot
                content=content,
                metadata=metadata
            )
            
            # Use custom chunk size if provided
            if custom_chunk_size and custom_chunk_size != Config.CHUNK_SIZE:
                text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=custom_chunk_size,
                    chunk_overlap=Config.CHUNK_OVERLAP,
                    separators=["\n\n", "\n", ". ", " ", ""]
                )
                chunks = self._create_chunks_with_splitter(content, filename, text_splitter)
            else:
                chunks = self._create_chunks(content, filename)
            
            chunk_count = len(chunks)
            
            # Generate and store embeddings
            embedding_success = True
            if chunks:
                embedding_success = self.embedding_manager.generate_and_store_embeddings(document_id, chunks)
                if not embedding_success:
                    logger.warning(f"Failed to generate embeddings for {filename}")
            
            # Extract and store contract information if requested
            contract_extracted = False
            contract_data = None
            if extract_contracts:
                contract_data = self._extract_contract_info(content, filename)
                if contract_data:
                    self.db_manager.insert_contract(document_id, contract_data)
                    contract_extracted = True
            
            logger.info(f"Successfully processed uploaded document {filename}: {chunk_count} chunks created")
            
            return {
                "success": True,
                "document_id": document_id,
                "filename": filename,
                "chunks_created": chunk_count,
                "contract_extracted": contract_extracted,
                "embedding_success": embedding_success,
                "content_length": len(content),
                "contract_data": contract_data
            }
            
        except Exception as e:
            logger.error(f"Error processing uploaded document {file_path}: {e}")
            return {
                "success": False,
                "error": str(e),
                "filename": filename or os.path.basename(file_path)
            }
    
    def _get_supported_files(self, folder_path: str) -> List[str]:
        """Get list of supported files in folder"""
        supported_extensions = {'.pdf', '.docx', '.txt'}
        supported_files = []
        
        for filename in os.listdir(folder_path):
            file_path = os.path.join(folder_path, filename)
            if os.path.isfile(file_path):
                _, ext = os.path.splitext(filename)
                if ext.lower() in supported_extensions:
                    supported_files.append(file_path)
        
        return supported_files
    
    def _extract_content(self, file_path: str, file_extension: str) -> tuple:
        """Extract content from different file types"""
        metadata = {
            "file_size": os.path.getsize(file_path),
            "created_at": datetime.fromtimestamp(os.path.getctime(file_path)).isoformat(),
            "modified_at": datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat(),
            "extraction_method": ""
        }
        
        if file_extension == '.pdf':
            return self._extract_from_pdf(file_path, metadata)
        elif file_extension == '.docx':
            return self._extract_from_docx(file_path, metadata)
        elif file_extension == '.txt':
            return self._extract_from_txt(file_path, metadata)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    def _extract_from_pdf(self, file_path: str, metadata: Dict) -> tuple:
        """Extract content from PDF files"""
        try:
            content = ""
            
            # First try to extract text directly
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        content += f"\n--- Page {page_num + 1} ---\n{page_text}"
                
                metadata["total_pages"] = len(pdf_reader.pages)
            
            # If no text extracted or very little text, try OCR
            if len(content.strip()) < 100:
                logger.info(f"Attempting OCR on {file_path} (little text extracted)")
                content = self._extract_with_ocr(file_path, metadata)
                metadata["extraction_method"] = "OCR"
            else:
                metadata["extraction_method"] = "direct_text"
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"Failed to extract from PDF {file_path}: {e}")
            # Fallback to OCR if PDF reading fails
            try:
                content = self._extract_with_ocr(file_path, metadata)
                metadata["extraction_method"] = "OCR_fallback"
                return content, metadata
            except Exception as ocr_error:
                logger.error(f"OCR fallback failed for {file_path}: {ocr_error}")
                return "", metadata
    
    def _extract_with_ocr(self, file_path: str, metadata: Dict) -> str:
        """Extract text using OCR for scanned PDFs"""
        try:
            # Convert PDF pages to images
            images = convert_from_path(file_path, dpi=300)
            content = ""
            
            for i, image in enumerate(images):
                # Perform OCR on each page
                page_text = pytesseract.image_to_string(image, lang='eng')
                if page_text.strip():
                    content += f"\n--- Page {i + 1} (OCR) ---\n{page_text}"
            
            metadata["ocr_pages"] = len(images)
            return content
            
        except Exception as e:
            logger.error(f"OCR extraction failed for {file_path}: {e}")
            return ""
    
    def _extract_from_docx(self, file_path: str, metadata: Dict) -> tuple:
        """Extract content from DOCX files"""
        try:
            doc = Document(file_path)
            content = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        content += row_text + "\n"
            
            metadata["extraction_method"] = "docx_parser"
            metadata["total_paragraphs"] = len(doc.paragraphs)
            metadata["total_tables"] = len(doc.tables)
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"Failed to extract from DOCX {file_path}: {e}")
            return "", metadata
    
    def _extract_from_txt(self, file_path: str, metadata: Dict) -> tuple:
        """Extract content from TXT files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            metadata["extraction_method"] = "text_file"
            metadata["character_count"] = len(content)
            metadata["line_count"] = content.count('\n')
            
            return content, metadata
            
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                    
                    metadata["extraction_method"] = f"text_file_{encoding}"
                    return content, metadata
                except UnicodeDecodeError:
                    continue
            
            logger.error(f"Failed to decode text file {file_path}")
            return "", metadata
    
    def _create_chunks(self, content: str, filename: str) -> List[str]:
        """Create text chunks from content"""
        try:
            # Create a document object for the text splitter
            doc = LangchainDocument(
                page_content=content,
                metadata={"source": filename}
            )
            
            # Split into chunks
            chunks = self.text_splitter.split_documents([doc])
            
            return [chunk.page_content for chunk in chunks]
            
        except Exception as e:
            logger.error(f"Failed to create chunks for {filename}: {e}")
            return []
    
    def _create_chunks_with_splitter(self, content: str, filename: str, splitter) -> List[str]:
        """Create text chunks from content using a custom splitter"""
        try:
            # Create a document object for the text splitter
            doc = LangchainDocument(
                page_content=content,
                metadata={"source": filename}
            )
            
            # Split into chunks using custom splitter
            chunks = splitter.split_documents([doc])
            
            return [chunk.page_content for chunk in chunks]
            
        except Exception as e:
            logger.error(f"Failed to create chunks for {filename} with custom splitter: {e}")
            return []
    
    def _extract_contract_info(self, content: str, filename: str) -> Optional[Dict[str, Any]]:
        """Extract structured contract information from content"""
        try:
            contract_data = {
                "contract_name": self._extract_contract_name(content, filename),
                "parties": self._extract_parties(content),
                "start_date": self._extract_date(content, "start"),
                "end_date": self._extract_date(content, "end"),
                "renewal_date": self._extract_date(content, "renewal"),
                "key_clauses": self._extract_key_clauses(content),
                "contact_info": self._extract_contact_info(content),
                "department": self._extract_department(content, filename)
            }
            
            # Only return if we found some meaningful information
            if any(contract_data.values()):
                return contract_data
            
            return None
            
        except Exception as e:
            logger.error(f"Failed to extract contract info from {filename}: {e}")
            return None
    
    def _extract_contract_name(self, content: str, filename: str) -> Optional[str]:
        """Extract contract name from content or filename"""
        # Try to extract from content first
        name_patterns = [
            r"Contract Title:?\s*(.+)",
            r"Agreement:\s*(.+)",
            r"CONTRACT:\s*(.+)",
            r"^(.+(?:Agreement|Contract))",
        ]
        
        for pattern in name_patterns:
            match = re.search(pattern, content, re.IGNORECASE | re.MULTILINE)
            if match:
                name = match.group(1).strip()
                if len(name) < 100:  # Reasonable length check
                    return name
        
        # Fallback to filename-based extraction
        base_name = os.path.splitext(filename)[0]
        return base_name.replace('_', ' ').title()
    
    def _extract_parties(self, content: str) -> List[str]:
        """Extract contract parties from content"""
        parties = []
        
        # Look for party patterns
        party_patterns = [
            r"Party 1:?\s*(.+?)(?:\n|Party 2)",
            r"Party 2:?\s*(.+?)(?:\n|$)",
            r"Parties:?\s*(.+?)(?:\n\n|\nCONTRACT)",
            r"Between:?\s*(.+?)\s*and\s*(.+?)(?:\n|$)",
        ]
        
        for pattern in party_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE | re.DOTALL)
            for match in matches:
                if isinstance(match, tuple):
                    parties.extend(match)
                else:
                    parties.append(match)
        
        # Look for company names in Inc., LLC, Corp format
        company_pattern = r'([A-Za-z\s]+(?:Inc\.|LLC|Corp\.?|Corporation|Ltd\.?|Limited))'
        companies = re.findall(company_pattern, content)
        parties.extend(companies)
        
        # Clean and deduplicate
        cleaned_parties = []
        for party in parties:
            cleaned = party.strip().replace('\n', ' ')
            if cleaned and len(cleaned) > 3 and cleaned not in cleaned_parties:
                cleaned_parties.append(cleaned)
        
        return cleaned_parties[:10]  # Limit to reasonable number
    
    def _extract_date(self, content: str, date_type: str) -> Optional[str]:
        """Extract dates from content"""
        date_patterns = {
            "start": [
                r"Start Date:?\s*(\d{4}-\d{2}-\d{2})",
                r"Effective Date:?\s*(\d{4}-\d{2}-\d{2})",
                r"Begin(?:ning)? Date:?\s*(\d{4}-\d{2}-\d{2})",
            ],
            "end": [
                r"End Date:?\s*(\d{4}-\d{2}-\d{2})",
                r"Expir(?:ation|y) Date:?\s*(\d{4}-\d{2}-\d{2})",
                r"Termination Date:?\s*(\d{4}-\d{2}-\d{2})",
            ],
            "renewal": [
                r"Renewal Date:?\s*(\d{4}-\d{2}-\d{2})",
                r"Review Date:?\s*(\d{4}-\d{2}-\d{2})",
            ]
        }
        
        patterns = date_patterns.get(date_type, [])
        for pattern in patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                return match.group(1)
        
        return None
    
    def _extract_key_clauses(self, content: str) -> List[str]:
        """Extract key clauses from content"""
        clauses = []
        
        # Look for numbered clauses
        clause_pattern = r"(\d+\.\s*.+?)(?=\n\d+\.|\n[A-Z]{2,}|$)"
        matches = re.findall(clause_pattern, content, re.DOTALL)
        
        for match in matches:
            clause = match.strip()
            if len(clause) > 20 and len(clause) < 500:  # Reasonable length
                clauses.append(clause)
        
        return clauses[:10]  # Limit number of clauses
    
    def _extract_contact_info(self, content: str) -> Dict[str, Any]:
        """Extract contact information from content"""
        contact_info = {}
        
        # Email pattern
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, content)
        if emails:
            contact_info["emails"] = list(set(emails))
        
        # Phone pattern
        phone_pattern = r'\+?1?[-.\s]?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}'
        phones = re.findall(phone_pattern, content)
        if phones:
            contact_info["phones"] = list(set(phones))
        
        # Address pattern (simplified)
        address_pattern = r'\d+\s+[A-Za-z\s,]+(?:Street|St|Avenue|Ave|Boulevard|Blvd|Drive|Dr|Way|Road|Rd),?\s+[A-Za-z\s,]+\d{5}'
        addresses = re.findall(address_pattern, content)
        if addresses:
            contact_info["addresses"] = list(set(addresses))
        
        return contact_info
    
    def _extract_department(self, content: str, filename: str) -> Optional[str]:
        """Extract department information"""
        departments = ["Legal", "IT", "Finance", "Procurement", "Operations", "HR"]
        
        for dept in departments:
            if dept.lower() in content.lower() or dept.lower() in filename.lower():
                return dept
        
        return None