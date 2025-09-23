"""
Synthetic dataset generator for CLM automation system.
Creates realistic contract documents with variations and conflicts.
"""

import os
import json
from datetime import datetime, timedelta
from typing import Dict, Any, List
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import random

class SyntheticDataGenerator:
    """Generates synthetic contract documents for testing"""
    
    def __init__(self, output_dir: str = "./documents"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        
        # Sample data for generating realistic contracts
        self.companies = [
            {"name": "TechCorp Solutions Inc.", "address": "123 Silicon Valley Blvd, San Jose, CA 95110", 
             "email": "contracts@techcorp.com", "phone": "+1-555-0123"},
            {"name": "GlobalTech Enterprises", "address": "456 Innovation Drive, Austin, TX 78701", 
             "email": "legal@globaltech.com", "phone": "+1-555-0124"},
            {"name": "DataSystems LLC", "address": "789 Tech Park Ave, Seattle, WA 98101", 
             "email": "admin@datasystems.com", "phone": "+1-555-0125"},
            {"name": "CloudVentures Corp", "address": "321 Digital Way, Denver, CO 80202", 
             "email": "partnerships@cloudventures.com", "phone": "+1-555-0126"},
            {"name": "InnovateLabs Inc.", "address": "654 Research Blvd, Boston, MA 02101", 
             "email": "contracts@innovatelabs.com", "phone": "+1-555-0127"}
        ]
        
        self.departments = ["Legal", "Procurement", "IT", "Finance", "Operations"]
        
        self.contract_types = [
            "Software License Agreement",
            "Service Level Agreement", 
            "Non-Disclosure Agreement",
            "Vendor Service Contract",
            "Maintenance Agreement"
        ]
    
    def generate_all_documents(self):
        """Generate the complete synthetic dataset"""
        print("Generating synthetic contract dataset...")
        
        # Generate different types of documents
        self._generate_pdf_contracts()
        self._generate_docx_contracts() 
        self._generate_txt_contracts()
        self._generate_unstructured_documents()
        
        print(f"Synthetic dataset generated in {self.output_dir}")
    
    def _generate_pdf_contracts(self):
        """Generate PDF contract documents"""
        contracts = [
            {
                "filename": "TechCorp_Software_License_2024.pdf",
                "parties": ["TechCorp Solutions Inc.", "GlobalTech Enterprises"],
                "contract_type": "Software License Agreement",
                "start_date": "2024-01-15",
                "end_date": "2025-01-14",
                "department": "IT",
                "amount": "$120,000"
            },
            {
                "filename": "DataSystems_SLA_2024.pdf", 
                "parties": ["DataSystems LLC", "CloudVentures Corp"],
                "contract_type": "Service Level Agreement",
                "start_date": "2024-03-01",
                "end_date": "2025-02-28", 
                "department": "Operations",
                "amount": "$85,000"
            },
            {
                "filename": "NDA_InnovateLabs_Scanned.pdf",
                "parties": ["InnovateLabs Inc.", "TechCorp Solutions Inc."],
                "contract_type": "Non-Disclosure Agreement", 
                "start_date": "2024-06-01",
                "end_date": "2026-05-31",
                "department": "Legal",
                "amount": "N/A"
            },
            {
                "filename": "GlobalTech_Maintenance_Agreement.pdf",
                "parties": ["GlobalTech Enterprises", "DataSystems LLC"],
                "contract_type": "Maintenance Agreement",
                "start_date": "2024-02-01", 
                "end_date": "2025-01-31",
                "department": "IT",
                "amount": "$45,000"
            },
            {
                "filename": "CloudVentures_Service_Contract_V2.pdf",
                "parties": ["CloudVentures Corp", "InnovateLabs Inc."],
                "contract_type": "Vendor Service Contract",
                "start_date": "2024-04-15",
                "end_date": "2024-10-14",  # This will be expiring soon
                "department": "Procurement", 
                "amount": "$67,500"
            }
        ]
        
        for contract in contracts:
            self._create_pdf_contract(contract)
    
    def _create_pdf_contract(self, contract_data: Dict[str, Any]):
        """Create a PDF contract document"""
        filepath = os.path.join(self.output_dir, contract_data["filename"])
        
        c = canvas.Canvas(filepath, pagesize=letter)
        width, height = letter
        
        # Title
        c.setFont("Helvetica-Bold", 18)
        c.drawString(50, height - 50, contract_data["contract_type"])
        
        # Contract details
        c.setFont("Helvetica", 12)
        y = height - 100
        
        lines = [
            f"Contract Date: {datetime.now().strftime('%B %d, %Y')}",
            f"",
            f"PARTIES:",
            f"Party 1: {contract_data['parties'][0]}",
            f"Party 2: {contract_data['parties'][1]}",
            f"",
            f"CONTRACT TERMS:",
            f"Start Date: {contract_data['start_date']}",
            f"End Date: {contract_data['end_date']}",
            f"Department: {contract_data['department']}",
            f"Total Amount: {contract_data['amount']}",
            f"",
            f"CLAUSES:",
            f"1. Payment Terms: Net 30 days from invoice date",
            f"2. Termination: Either party may terminate with 30 days notice",
            f"3. Renewal: Automatic renewal unless terminated",
            f"4. Confidentiality: All information remains confidential",
            f"5. Force Majeure: Standard force majeure provisions apply",
            f"",
            f"CONTACT INFORMATION:",
            f"{self._get_company_contact(contract_data['parties'][0])}",
            f"{self._get_company_contact(contract_data['parties'][1])}",
            f"",
            f"This contract is governed by the laws of California.",
            f"",
            f"Signatures:",
            f"_____________________    _____________________",
            f"Party 1 Signature        Party 2 Signature"
        ]
        
        for line in lines:
            c.drawString(50, y, line)
            y -= 20
        
        c.save()
        print(f"Created PDF: {contract_data['filename']}")
    
    def _generate_docx_contracts(self):
        """Generate DOCX contract documents"""
        contracts = [
            {
                "filename": "TechCorp_Amendment_2024.docx",
                "parties": ["TechCorp Solutions Inc.", "GlobalTech Enterprises"],
                "contract_type": "Software License Amendment",
                "original_contract": "TechCorp_Software_License_2024.pdf",
                "amendment_date": "2024-06-15",
                "changes": "Extended support period and added new modules"
            },
            {
                "filename": "DataSystems_Draft_Contract.docx", 
                "parties": ["DataSystems LLC", "InnovateLabs Inc."],
                "contract_type": "Service Agreement Draft",
                "start_date": "2024-08-01",
                "end_date": "2025-07-31", 
                "status": "DRAFT - Under Review"
            },
            {
                "filename": "NDA_CloudVentures_Updated.docx",
                "parties": ["CloudVentures Corp", "TechCorp Solutions Inc."],
                "contract_type": "Non-Disclosure Agreement", 
                "start_date": "2024-07-01",
                "end_date": "2026-06-30",
                "department": "Legal",
                "conflicts": "Different address for TechCorp than in other contracts"
            },
            {
                "filename": "GlobalTech_Service_Extension.docx",
                "parties": ["GlobalTech Enterprises", "CloudVentures Corp"],
                "contract_type": "Service Extension Agreement",
                "start_date": "2024-05-01",
                "end_date": "2024-11-30",
                "department": "Operations"
            }
        ]
        
        for contract in contracts:
            self._create_docx_contract(contract)
    
    def _create_docx_contract(self, contract_data: Dict[str, Any]):
        """Create a DOCX contract document"""
        doc = Document()
        
        # Title
        title = doc.add_heading(contract_data["contract_type"], 0)
        
        # Add contract content
        doc.add_paragraph(f"Document Date: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph("")
        
        doc.add_heading('Parties:', level=1)
        for party in contract_data["parties"]:
            doc.add_paragraph(f"• {party}")
        
        if "original_contract" in contract_data:
            doc.add_heading('Amendment Details:', level=1)
            doc.add_paragraph(f"Original Contract: {contract_data['original_contract']}")
            doc.add_paragraph(f"Amendment Date: {contract_data['amendment_date']}")
            doc.add_paragraph(f"Changes: {contract_data['changes']}")
        
        if "start_date" in contract_data:
            doc.add_heading('Terms:', level=1)
            doc.add_paragraph(f"Start Date: {contract_data['start_date']}")
            doc.add_paragraph(f"End Date: {contract_data['end_date']}")
            
        if "department" in contract_data:
            doc.add_paragraph(f"Responsible Department: {contract_data['department']}")
            
        if "status" in contract_data:
            doc.add_paragraph(f"Status: {contract_data['status']}")
            
        if "conflicts" in contract_data:
            doc.add_paragraph(f"Note: {contract_data['conflicts']}")
        
        # Contact information (with intentional conflicts)
        doc.add_heading('Contact Information:', level=1)
        for party in contract_data["parties"]:
            contact = self._get_company_contact(party)
            # Introduce conflicts for testing
            if "CloudVentures" in party and "NDA" in contract_data["filename"]:
                contact = contact.replace("321 Digital Way", "999 Conflict Street")
            doc.add_paragraph(contact)
        
        filepath = os.path.join(self.output_dir, contract_data["filename"])
        doc.save(filepath)
        print(f"Created DOCX: {contract_data['filename']}")
    
    def _generate_txt_contracts(self):
        """Generate TXT contract summaries and correspondence"""
        
        # Contract summary
        summary_content = """
CONTRACT SUMMARY REPORT
Generated: {date}

Active Contracts Summary:
========================

1. TechCorp Software License Agreement
   - Parties: TechCorp Solutions Inc., GlobalTech Enterprises  
   - Term: 2024-01-15 to 2025-01-14
   - Value: $120,000
   - Department: IT
   - Status: Active
   - Contact: contracts@techcorp.com, +1-555-0123

2. DataSystems Service Level Agreement
   - Parties: DataSystems LLC, CloudVentures Corp
   - Term: 2024-03-01 to 2025-02-28
   - Value: $85,000
   - Department: Operations
   - Status: Active
   - Contact: admin@datasystems.com, +1-555-0125

3. InnovateLabs Non-Disclosure Agreement
   - Parties: InnovateLabs Inc., TechCorp Solutions Inc.
   - Term: 2024-06-01 to 2026-05-31
   - Value: N/A
   - Department: Legal
   - Status: Active
   - Contact: contracts@innovatelabs.com, +1-555-0127

EXPIRING SOON:
- CloudVentures Service Contract expires 2024-10-14 (URGENT)

CONFLICTS DETECTED:
- TechCorp address mismatch between contracts
- Different contact info for CloudVentures in NDA vs SLA
        """.format(date=datetime.now().strftime('%B %d, %Y'))
        
        with open(os.path.join(self.output_dir, "contract_summary_2024.txt"), 'w') as f:
            f.write(summary_content)
        
        # Email correspondence
        email_content = """
From: legal@techcorp.com
To: contracts@globaltech.com  
Subject: Contract Renewal Discussion - Software License Agreement
Date: {date}

Dear GlobalTech Legal Team,

I hope this email finds you well. I'm reaching out regarding our Software License Agreement that expires on January 14, 2025.

Key Points for Renewal Discussion:
- Current contract value: $120,000 annually
- Excellent service delivery over the past year
- Request for 5% increase to $126,000 for next term
- Additional modules to be included in renewal

We would like to schedule a meeting to discuss renewal terms. Please note that our new office address is 123 Silicon Valley Blvd, San Jose, CA 95110 (updated from previous communications).

Contact Information:
Primary: contracts@techcorp.com
Phone: +1-555-0123  
Legal Dept: Sarah Johnson, Director of Legal Affairs

Looking forward to continuing our partnership.

Best regards,
TechCorp Legal Team

---
CONFIDENTIALITY NOTICE: This email contains confidential information intended solely for the addressee.
        """.format(date=datetime.now().strftime('%B %d, %Y'))
        
        with open(os.path.join(self.output_dir, "email_correspondence_contract_renewal.txt"), 'w') as f:
            f.write(email_content)
            
        # Contract amendment notes
        amendment_content = """
AMENDMENT TRACKING NOTES
Contract: TechCorp_Software_License_2024.pdf
Last Updated: {date}

CHANGE LOG:
===========

Version 1.0 (Original): January 15, 2024
- Initial contract signed
- Term: 12 months  
- Value: $120,000
- Standard terms and conditions

Version 1.1 (Amendment 1): June 15, 2024  
- Added cloud storage module (+$15,000)
- Extended support hours to 24/7
- Updated contact: Changed primary contact from john@techcorp.com to contracts@techcorp.com
- No change to end date

PENDING CHANGES:
- Renewal negotiation in progress
- Proposed rate increase to $126,000
- Additional security modules requested
- New 2-year term proposed

NOTES:
- GlobalTech has been reliable partner
- Payment history: Excellent (no late payments)
- Support tickets resolved promptly
- Consider preferred vendor status for renewal

ACTION ITEMS:
[ ] Schedule renewal meeting by September 30
[ ] Prepare renewal proposal with updated terms
[ ] Legal review of new clauses
[ ] Get executive approval for pricing changes
        """.format(date=datetime.now().strftime('%B %d, %Y'))
        
        with open(os.path.join(self.output_dir, "amendment_notes_techcorp.txt"), 'w') as f:
            f.write(amendment_content)
            
        print("Created TXT files: contract_summary_2024.txt, email_correspondence_contract_renewal.txt, amendment_notes_techcorp.txt")
    
    def _generate_unstructured_documents(self):
        """Generate unstructured meeting notes and informal documents"""
        
        # Meeting notes
        meeting_notes = """
        MEETING NOTES - CONTRACT REVIEW SESSION
        Date: {date}
        Attendees: Sarah (Legal), Mike (Procurement), Lisa (Finance), Tom (IT)
        
        Discussion Points:
        - Reviewed Q3 contract performance
        - DataSystems contract performing well, recommend renewal
        - Issue with CloudVentures billing - address discrepancies
        - TechCorp wants to add new modules, legal review needed
        
        ACTION ITEMS:
        Sarah: Review TechCorp amendment by Friday
        Mike: Negotiate better terms with CloudVentures  
        Lisa: Prepare budget for contract renewals
        Tom: Technical evaluation of new TechCorp modules
        
        CONCERNS RAISED:
        1. CloudVentures has different contact info in various documents
           - SLA shows: partnerships@cloudventures.com
           - NDA shows: legal@cloudventures.com  
           - Need to clarify correct contact
        
        2. Multiple contract versions for same parties causing confusion
           - TechCorp has original + amendment + renewal draft
           - Version control system needed
        
        3. Expiring contracts need attention:
           - CloudVentures service contract expires soon (Oct 2024)
           - Need renewal discussion ASAP
        
        NEXT MEETING: Next Friday 2PM
        Location: Conference Room B
        Agenda: Contract renewal priorities
        """.format(date=datetime.now().strftime('%B %d, %Y'))
        
        with open(os.path.join(self.output_dir, "meeting_notes_contract_review.txt"), 'w') as f:
            f.write(meeting_notes)
        
        # Informal contract discussion notes  
        informal_notes = """
        Contract Discussion - Slack Thread Archive
        Channel: #legal-contracts
        Date Range: Last 30 days
        
        @sarah.legal: Hey team, quick update on contracts
        - TechCorp renewal looking good
        - They want 2-year term this time
        - Pricing increase reasonable at 5%
        
        @mike.procurement: Sounds good. What about DataSystems?
        Their SLA has been solid, no issues
        
        @lisa.finance: Budget wise we're okay with both renewals
        CloudVentures is the one I'm worried about - they've had billing errors
        
        @tom.it: From technical side:
        - TechCorp modules work great, approve new ones
        - DataSystems uptime has been 99.8% - excellent  
        - CloudVentures had that outage last month...
        
        @sarah.legal: Speaking of CloudVentures, found discrepancy in their contact info
        Some docs have different email/address than others
        Need to verify which is current
        
        @mike.procurement: I'll reach out to them directly
        Their main contact is usually responsive
        
        @lisa.finance: FYI - CloudVentures contract expires in 2 months
        If we're renewing need to start talks soon
        
        @sarah.legal: Good catch! Adding to urgent list
        Will send renewal notice this week
        
        @tom.it: Also heads up - InnovateLabs asked about early renewal
        They're happy with current NDA terms
        
        @mike.procurement: That's great, one less thing to negotiate
        Simple renewal should be quick
        
        @sarah.legal: Agreed. Summary of priorities:
        1. CloudVentures renewal (urgent - expires soon)
        2. TechCorp 2-year renewal (in progress)  
        3. DataSystems renewal (routine)
        4. InnovateLabs NDA renewal (simple)
        
        All: Sounds like a plan! 👍
        """.format(date=datetime.now().strftime('%B %d, %Y'))
        
        with open(os.path.join(self.output_dir, "informal_contract_discussions.txt"), 'w') as f:
            f.write(informal_notes)
            
        print("Created unstructured files: meeting_notes_contract_review.txt, informal_contract_discussions.txt")
    
    def _get_company_contact(self, company_name: str) -> str:
        """Get contact information for a company"""
        for company in self.companies:
            if company["name"] in company_name:
                return f"{company['name']}\n{company['address']}\nEmail: {company['email']}\nPhone: {company['phone']}"
        return f"{company_name}\nContact information not available"


if __name__ == "__main__":
    generator = SyntheticDataGenerator()
    generator.generate_all_documents()