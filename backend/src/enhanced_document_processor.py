"""
Enhanced document processing module for CLM automation system.
Handles document ingestion with OpenAI Vision API support and large document processing.
"""

import os
import logging
import re
import base64
import io
from typing import Dict, Any, List, Optional, Tuple, Generator
from datetime import datetime
import json
import asyncio
from concurrent.futures import ThreadPoolExecutor
import tempfile

# File processing imports
import PyPDF2
from docx import Document
import pytesseract
from pdf2image import convert_from_path
from PIL import Image
import fitz  # PyMuPDF for better PDF handling

# LangChain imports
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangchainDocument

# OpenAI imports
from openai import OpenAI

from src.config import Config
from src.database import DatabaseManager
from src.embeddings import EmbeddingManager

logger = logging.getLogger(__name__)

class EnhancedDocumentProcessor:
    """Enhanced document processor with Vision API and large document support"""
    
    def __init__(self, db_manager: DatabaseManager):
        self.db_manager = db_manager
        self.embedding_manager = EmbeddingManager(db_manager)
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=Config.CHUNK_SIZE,
            chunk_overlap=Config.CHUNK_OVERLAP,
            separators=["\\n\\n", "\\n", ". ", " ", ""]
        )
        self.openai_client = OpenAI(api_key=Config.OPENAI_API_KEY)
        
        # Configuration for large document processing
        self.max_pdf_pages_direct = 50  # Process directly up to this many pages
        self.batch_size = 10  # Process pages in batches
        self.max_image_size = (1024, 1024)  # Resize images to this max size
        self.vision_enabled = True
        
        # Progress tracking
        self.current_progress = {}
    
    async def process_document_async(self, file_path: str, filename: str = None, 
                                   extract_contracts: bool = True, 
                                   custom_chunk_size: int = None,
                                   metadata: Dict[str, Any] = None,
                                   use_vision: bool = True,
                                   progress_callback=None) -> Dict[str, Any]:
        """Asynchronously process a single document with enhanced features"""
        try:
            if filename is None:
                filename = os.path.basename(file_path)
            
            logger.info(f"Starting enhanced processing of: {filename}")
            
            # Initialize progress tracking
            progress_id = f"doc_{datetime.now().timestamp()}"
            self.current_progress[progress_id] = {
                "filename": filename,
                "status": "starting",
                "progress": 0,
                "total_steps": 6,
                "current_step": "initialization"
            }
            
            if progress_callback:
                await progress_callback(progress_id, self.current_progress[progress_id])
            
            file_extension = os.path.splitext(filename)[1].lower()
            
            # Step 1: Extract content with enhanced methods
            self._update_progress(progress_id, 1, "extracting_content", progress_callback)
            content, file_metadata, visual_content = await self._extract_content_enhanced(
                file_path, file_extension, use_vision
            )
            
            if not content.strip() and not visual_content:
                logger.warning(f"No content extracted from {filename}")
                return self._create_error_result(filename, "No content could be extracted from the document")
            
            # Step 2: Process and merge metadata
            self._update_progress(progress_id, 2, "processing_metadata", progress_callback)
            combined_metadata = self._merge_metadata(file_metadata, metadata, visual_content)
            
            # Step 3: Store document
            self._update_progress(progress_id, 3, "storing_document", progress_callback)
            document_id = self.db_manager.insert_document(
                filename=filename,
                file_type=file_extension[1:],
                content=content,
                metadata=combined_metadata
            )
            
            # Step 4: Create enhanced chunks
            self._update_progress(progress_id, 4, "creating_chunks", progress_callback)
            chunks = await self._create_enhanced_chunks(
                content, filename, visual_content, custom_chunk_size
            )
            
            # Step 5: Generate embeddings
            self._update_progress(progress_id, 5, "generating_embeddings", progress_callback)
            embedding_success = await self._generate_embeddings_async(document_id, chunks)
            
            # Step 6: Extract contracts
            self._update_progress(progress_id, 6, "extracting_contracts", progress_callback)
            contract_extracted, contract_data = await self._extract_contracts_async(
                document_id, content, filename, visual_content, extract_contracts
            )
            
            # Complete processing
            self._update_progress(progress_id, 6, "completed", progress_callback)
            
            result = {
                "success": True,
                "document_id": document_id,
                "filename": filename,
                "chunks_created": len(chunks),
                "contract_extracted": contract_extracted,
                "embedding_success": embedding_success,
                "content_length": len(content),
                "contract_data": contract_data,
                "visual_content_found": len(visual_content) > 0,
                "visual_elements": len(visual_content),
                "progress_id": progress_id
            }
            
            logger.info(f"Successfully processed {filename}: {len(chunks)} chunks, "
                       f"{len(visual_content)} visual elements")
            return result
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            return self._create_error_result(filename or os.path.basename(file_path), str(e))
    
    async def _extract_content_enhanced(self, file_path: str, file_extension: str, 
                                      use_vision: bool) -> Tuple[str, Dict, List[Dict]]:
        """Enhanced content extraction with vision support"""
        metadata = {
            "file_size": os.path.getsize(file_path),
            "created_at": datetime.fromtimestamp(os.path.getctime(file_path)).isoformat(),
            "modified_at": datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat(),
            "extraction_method": "",
            "processing_time": datetime.now().isoformat()
        }
        
        visual_content = []
        
        if file_extension == '.pdf':
            content, metadata, visual_content = await self._extract_from_pdf_enhanced(
                file_path, metadata, use_vision
            )
        elif file_extension == '.docx':
            content, metadata, visual_content = await self._extract_from_docx_enhanced(
                file_path, metadata, use_vision
            )
        elif file_extension == '.txt':
            content, metadata = self._extract_from_txt(file_path, metadata)
        elif file_extension in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff']:
            # Direct image processing
            content, metadata, visual_content = await self._extract_from_image(
                file_path, metadata, use_vision
            )
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        return content, metadata, visual_content
    
    async def _extract_from_pdf_enhanced(self, file_path: str, metadata: Dict, 
                                       use_vision: bool) -> Tuple[str, Dict, List[Dict]]:
        """Enhanced PDF extraction with Vision API support"""
        content = ""
        visual_content = []
        
        try:
            # Use PyMuPDF for better handling
            doc = fitz.open(file_path)
            total_pages = doc.page_count
            metadata["total_pages"] = total_pages
            
            logger.info(f"Processing PDF with {total_pages} pages")
            
            if total_pages > self.max_pdf_pages_direct:
                # Process large PDFs in batches
                content, visual_content = await self._process_large_pdf(
                    doc, metadata, use_vision
                )
            else:
                # Process smaller PDFs normally
                content, visual_content = await self._process_standard_pdf(
                    doc, metadata, use_vision
                )
            
            doc.close()
            
            # If still no content, try OCR fallback
            if len(content.strip()) < 100:
                logger.info(f"Attempting OCR fallback for {file_path}")
                ocr_content = await self._extract_with_ocr_enhanced(file_path, metadata)
                if ocr_content:
                    content = ocr_content
                    metadata["extraction_method"] += "_ocr_fallback"
            
            return content, metadata, visual_content
            
        except Exception as e:
            logger.error(f"Enhanced PDF extraction failed for {file_path}: {e}")
            # Fallback to basic extraction
            return await self._fallback_pdf_extraction(file_path, metadata)
    
    async def _process_large_pdf(self, doc, metadata: Dict, use_vision: bool) -> Tuple[str, List[Dict]]:
        """Process large PDFs in batches to avoid memory issues"""
        content = ""
        visual_content = []
        total_pages = doc.page_count
        
        logger.info(f"Processing large PDF with {total_pages} pages in batches")
        metadata["extraction_method"] = "large_pdf_batched"
        
        # Process pages in batches
        for batch_start in range(0, total_pages, self.batch_size):
            batch_end = min(batch_start + self.batch_size, total_pages)
            
            logger.info(f"Processing pages {batch_start + 1} to {batch_end}")
            
            batch_content = ""
            batch_visual = []
            
            for page_num in range(batch_start, batch_end):
                page = doc[page_num]
                
                # Extract text
                page_text = page.get_text()
                if page_text.strip():
                    batch_content += f"\\n--- Page {page_num + 1} ---\\n{page_text}"
                
                # Extract images if vision is enabled
                if use_vision and self.vision_enabled:
                    page_images = await self._extract_images_from_page(page, page_num + 1)
                    batch_visual.extend(page_images)
            
            content += batch_content
            visual_content.extend(batch_visual)
            
            # Process visual content with Vision API in batches
            if use_vision and batch_visual:
                await self._process_visual_content_batch(batch_visual)
        
        metadata["processed_pages"] = total_pages
        metadata["batch_size"] = self.batch_size
        
        return content, visual_content
    
    async def _process_standard_pdf(self, doc, metadata: Dict, use_vision: bool) -> Tuple[str, List[Dict]]:
        """Process standard-sized PDFs"""
        content = ""
        visual_content = []
        
        metadata["extraction_method"] = "standard_pdf"
        
        for page_num in range(doc.page_count):
            page = doc[page_num]
            
            # Extract text
            page_text = page.get_text()
            if page_text.strip():
                content += f"\\n--- Page {page_num + 1} ---\\n{page_text}"
            
            # Extract images if vision is enabled
            if use_vision and self.vision_enabled:
                page_images = await self._extract_images_from_page(page, page_num + 1)
                visual_content.extend(page_images)
        
        # Process all visual content
        if use_vision and visual_content:
            await self._process_visual_content_batch(visual_content)
        
        return content, visual_content
    
    async def _extract_images_from_page(self, page, page_num: int) -> List[Dict]:
        """Extract images from a PDF page"""
        images = []
        
        try:
            image_list = page.get_images()
            
            for img_index, img in enumerate(image_list):
                xref = img[0]
                pix = fitz.Pixmap(page.parent, xref)
                
                if pix.n - pix.alpha < 4:  # GRAY or RGB
                    # Convert to PIL Image
                    img_data = pix.tobytes("ppm")
                    pil_image = Image.open(io.BytesIO(img_data))
                    
                    # Resize if too large
                    if pil_image.size[0] > self.max_image_size[0] or pil_image.size[1] > self.max_image_size[1]:
                        pil_image.thumbnail(self.max_image_size, Image.Resampling.LANCZOS)
                    
                    # Convert to base64 for Vision API
                    buffered = io.BytesIO()
                    pil_image.save(buffered, format="PNG")
                    img_base64 = base64.b64encode(buffered.getvalue()).decode()
                    
                    images.append({
                        "page": page_num,
                        "image_index": img_index,
                        "base64": img_base64,
                        "size": pil_image.size,
                        "extracted_text": "",  # Will be filled by Vision API
                        "description": ""  # Will be filled by Vision API
                    })
                
                pix = None  # Free memory
                
        except Exception as e:
            logger.warning(f"Failed to extract images from page {page_num}: {e}")
        
        return images
    
    async def _process_visual_content_batch(self, visual_content: List[Dict]):
        """Process visual content using OpenAI Vision API in batches"""
        if not self.vision_enabled or not visual_content:
            return
        
        logger.info(f"Processing {len(visual_content)} visual elements with Vision API")
        
        # Process images in smaller batches to avoid rate limits
        batch_size = 5
        for i in range(0, len(visual_content), batch_size):
            batch = visual_content[i:i + batch_size]
            
            # Process batch concurrently
            tasks = [self._analyze_image_with_vision(img_data) for img_data in batch]
            results = await asyncio.gather(*tasks, return_exceptions=True)
            
            # Update the visual content with results
            for j, result in enumerate(results):
                if isinstance(result, Exception):
                    logger.warning(f"Vision API failed for image {i+j}: {result}")
                else:
                    batch[j].update(result)
            
            # Small delay to respect rate limits
            await asyncio.sleep(0.5)
    
    async def _analyze_image_with_vision(self, img_data: Dict) -> Dict:
        """Analyze a single image using OpenAI Vision API"""
        try:
            response = await asyncio.to_thread(
                self.openai_client.chat.completions.create,
                model="gpt-4-vision-preview",
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": "Analyze this image from a legal document. Extract any text, describe the content, and identify any important legal elements like signatures, seals, charts, or diagrams. Be concise but thorough."
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/png;base64,{img_data['base64']}"
                                }
                            }
                        ]
                    }
                ],
                max_tokens=500
            )
            
            analysis = response.choices[0].message.content
            
            # Parse the response to extract text and description
            # This is a simplified parser - you might want to make it more sophisticated
            extracted_text = ""
            description = analysis
            
            if "TEXT:" in analysis.upper():
                parts = analysis.split("TEXT:")
                if len(parts) > 1:
                    extracted_text = parts[1].split("DESCRIPTION:")[0].strip()
                    if "DESCRIPTION:" in analysis.upper():
                        description = analysis.split("DESCRIPTION:")[1].strip()
            
            return {
                "extracted_text": extracted_text,
                "description": description,
                "vision_analysis": analysis
            }
            
        except Exception as e:
            logger.error(f"Vision API analysis failed: {e}")
            return {
                "extracted_text": "",
                "description": f"Vision analysis failed: {str(e)}",
                "vision_analysis": ""
            }
    
    async def _extract_from_docx_enhanced(self, file_path: str, metadata: Dict, 
                                        use_vision: bool) -> Tuple[str, Dict, List[Dict]]:
        """Enhanced DOCX extraction with image support"""
        content = ""
        visual_content = []
        
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        content += row_text + "\\n"
            
            # Extract images if vision is enabled
            if use_vision and self.vision_enabled:
                # DOCX image extraction would require python-docx2txt or similar
                # For now, we'll note that images are present
                logger.info("DOCX image extraction not yet implemented")
            
            metadata["extraction_method"] = "enhanced_docx_parser"
            metadata["total_paragraphs"] = len(doc.paragraphs)
            metadata["total_tables"] = len(doc.tables)
            
            return content, metadata, visual_content
            
        except Exception as e:
            logger.error(f"Enhanced DOCX extraction failed for {file_path}: {e}")
            return "", metadata, []
    
    async def _extract_from_image(self, file_path: str, metadata: Dict, 
                                use_vision: bool) -> Tuple[str, Dict, List[Dict]]:
        """Extract content from image files using Vision API"""
        content = ""
        visual_content = []
        
        try:
            # Load and resize image
            pil_image = Image.open(file_path)
            original_size = pil_image.size
            
            if pil_image.size[0] > self.max_image_size[0] or pil_image.size[1] > self.max_image_size[1]:
                pil_image.thumbnail(self.max_image_size, Image.Resampling.LANCZOS)
            
            # Convert to base64
            buffered = io.BytesIO()
            pil_image.save(buffered, format="PNG")
            img_base64 = base64.b64encode(buffered.getvalue()).decode()
            
            visual_content.append({
                "page": 1,
                "image_index": 0,
                "base64": img_base64,
                "size": pil_image.size,
                "original_size": original_size,
                "extracted_text": "",
                "description": ""
            })
            
            # Process with Vision API if enabled
            if use_vision and self.vision_enabled:
                await self._process_visual_content_batch(visual_content)
                # Use extracted text as main content
                content = visual_content[0].get("extracted_text", "")
            
            # Fallback to OCR if no Vision API or no text extracted
            if not content.strip():
                logger.info("Using OCR fallback for image")
                content = pytesseract.image_to_string(pil_image)
                metadata["extraction_method"] = "ocr_fallback"
            else:
                metadata["extraction_method"] = "vision_api"
            
            metadata["original_size"] = original_size
            metadata["processed_size"] = pil_image.size
            
            return content, metadata, visual_content
            
        except Exception as e:
            logger.error(f"Image extraction failed for {file_path}: {e}")
            return "", metadata, []
    
    async def _extract_with_ocr_enhanced(self, file_path: str, metadata: Dict) -> str:
        """Enhanced OCR extraction with better error handling"""
        try:
            images = convert_from_path(file_path, dpi=300)
            content = ""
            
            for i, image in enumerate(images):
                # Resize if too large
                if image.size[0] > self.max_image_size[0] or image.size[1] > self.max_image_size[1]:
                    image.thumbnail(self.max_image_size, Image.Resampling.LANCZOS)
                
                # Perform OCR
                page_text = await asyncio.to_thread(
                    pytesseract.image_to_string, image, lang='eng'
                )
                if page_text.strip():
                    content += f"\\n--- Page {i + 1} (Enhanced OCR) ---\\n{page_text}"
            
            metadata["ocr_pages"] = len(images)
            metadata["ocr_method"] = "enhanced"
            return content
            
        except Exception as e:
            logger.error(f"Enhanced OCR extraction failed: {e}")
            return ""
    
    async def _create_enhanced_chunks(self, content: str, filename: str, 
                                    visual_content: List[Dict], 
                                    custom_chunk_size: int = None) -> List[Dict]:
        """Create enhanced chunks that include visual content references"""
        try:
            # Use custom chunk size if provided
            if custom_chunk_size and custom_chunk_size != Config.CHUNK_SIZE:
                text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=custom_chunk_size,
                    chunk_overlap=Config.CHUNK_OVERLAP,
                    separators=["\\n\\n", "\\n", ". ", " ", ""]
                )
            else:
                text_splitter = self.text_splitter
            
            # Create enhanced content by integrating visual descriptions
            enhanced_content = content
            
            # Add visual content descriptions to relevant sections
            for visual in visual_content:
                if visual.get("description"):
                    page_marker = f"--- Page {visual['page']} ---"
                    if page_marker in enhanced_content:
                        # Insert visual description after page marker
                        insertion_point = enhanced_content.find(page_marker) + len(page_marker)
                        visual_text = f"\\n[VISUAL CONTENT: {visual['description']}]\\n"
                        if visual.get("extracted_text"):
                            visual_text += f"[EXTRACTED TEXT: {visual['extracted_text']}]\\n"
                        
                        enhanced_content = (enhanced_content[:insertion_point] + 
                                          visual_text + 
                                          enhanced_content[insertion_point:])
            
            # Create document for splitting
            doc = LangchainDocument(
                page_content=enhanced_content,
                metadata={"source": filename, "has_visual_content": len(visual_content) > 0}
            )
            
            # Split into chunks
            chunks = text_splitter.split_documents([doc])
            
            # Convert to enhanced chunk format
            enhanced_chunks = []
            for i, chunk in enumerate(chunks):
                enhanced_chunk = {
                    "content": chunk.page_content,
                    "metadata": chunk.metadata,
                    "chunk_index": i,
                    "visual_references": []
                }
                
                # Find visual content that belongs to this chunk
                for visual in visual_content:
                    if (visual.get("description") and 
                        visual["description"][:50] in chunk.page_content):
                        enhanced_chunk["visual_references"].append({
                            "page": visual["page"],
                            "description": visual["description"],
                            "extracted_text": visual.get("extracted_text", "")
                        })
                
                enhanced_chunks.append(enhanced_chunk)
            
            return enhanced_chunks
            
        except Exception as e:
            logger.error(f"Failed to create enhanced chunks for {filename}: {e}")
            return []
    
    async def _generate_embeddings_async(self, document_id: str, chunks: List[Dict]) -> bool:
        """Generate embeddings asynchronously"""
        try:
            if not chunks:
                return False
            
            # Convert enhanced chunks to simple text for embedding
            text_chunks = [chunk["content"] for chunk in chunks]
            
            # Generate embeddings
            success = await asyncio.to_thread(
                self.embedding_manager.generate_and_store_embeddings,
                document_id, text_chunks
            )
            
            return success
            
        except Exception as e:
            logger.error(f"Failed to generate embeddings: {e}")
            return False
    
    async def _extract_contracts_async(self, document_id: str, content: str, 
                                     filename: str, visual_content: List[Dict], 
                                     extract_contracts: bool) -> Tuple[bool, Optional[Dict]]:
        """Extract contract information asynchronously"""
        try:
            if not extract_contracts:
                return False, None
            
            # Enhanced contract extraction including visual content
            enhanced_content = content
            
            # Add visual text to content for contract extraction
            for visual in visual_content:
                if visual.get("extracted_text"):
                    enhanced_content += f"\\n{visual['extracted_text']}"
            
            contract_data = await asyncio.to_thread(
                self._extract_contract_info_enhanced, 
                enhanced_content, filename, visual_content
            )
            
            if contract_data:
                self.db_manager.insert_contract(document_id, contract_data)
                return True, contract_data
            
            return False, None
            
        except Exception as e:
            logger.error(f"Failed to extract contracts: {e}")
            return False, None
    
    def _extract_contract_info_enhanced(self, content: str, filename: str, 
                                      visual_content: List[Dict]) -> Optional[Dict[str, Any]]:
        """Enhanced contract information extraction"""
        try:
            # Basic contract extraction (reuse existing logic)
            contract_data = {
                "contract_name": self._extract_contract_name(content, filename),
                "parties": self._extract_parties(content),
                "start_date": self._extract_date(content, "start"),
                "end_date": self._extract_date(content, "end"),
                "renewal_date": self._extract_date(content, "renewal"),
                "key_clauses": self._extract_key_clauses(content),
                "contact_info": self._extract_contact_info(content),
                "department": self._extract_department(content, filename),
                "visual_elements": len(visual_content),
                "has_signatures": self._detect_signatures(visual_content),
                "has_seals": self._detect_seals(visual_content)
            }
            
            # Only return if we found meaningful information
            if any(v for k, v in contract_data.items() if k not in ['visual_elements', 'has_signatures', 'has_seals']):
                return contract_data
            
            return None
            
        except Exception as e:
            logger.error(f"Enhanced contract extraction failed: {e}")
            return None
    
    def _detect_signatures(self, visual_content: List[Dict]) -> bool:
        """Detect signatures in visual content"""
        for visual in visual_content:
            description = visual.get("description", "").lower()
            if any(term in description for term in ["signature", "signed", "sign"]):
                return True
        return False
    
    def _detect_seals(self, visual_content: List[Dict]) -> bool:
        """Detect official seals in visual content"""
        for visual in visual_content:
            description = visual.get("description", "").lower()
            if any(term in description for term in ["seal", "stamp", "official", "notary"]):
                return True
        return False
    
    def _merge_metadata(self, file_metadata: Dict, custom_metadata: Dict, 
                       visual_content: List[Dict]) -> Dict:
        """Merge all metadata sources"""
        combined = file_metadata.copy()
        
        if custom_metadata:
            combined.update(custom_metadata)
        
        # Add visual content metadata
        combined["visual_content_count"] = len(visual_content)
        combined["has_visual_content"] = len(visual_content) > 0
        
        if visual_content:
            combined["visual_pages"] = list(set(v["page"] for v in visual_content))
        
        return combined
    
    def _update_progress(self, progress_id: str, step: int, status: str, callback):
        """Update progress tracking"""
        if progress_id in self.current_progress:
            self.current_progress[progress_id].update({
                "progress": (step / self.current_progress[progress_id]["total_steps"]) * 100,
                "current_step": status,
                "status": status
            })
            
            if callback:
                asyncio.create_task(callback(progress_id, self.current_progress[progress_id]))
    
    def _create_error_result(self, filename: str, error: str) -> Dict[str, Any]:
        """Create standardized error result"""
        return {
            "success": False,
            "error": error,
            "filename": filename,
            "chunks_created": 0,
            "contract_extracted": False,
            "embedding_success": False,
            "visual_content_found": False
        }
    
    async def _fallback_pdf_extraction(self, file_path: str, metadata: Dict) -> Tuple[str, Dict, List[Dict]]:
        """Fallback PDF extraction using original method"""
        try:
            content = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        content += f"\\n--- Page {page_num + 1} ---\\n{page_text}"
                
                metadata["total_pages"] = len(pdf_reader.pages)
                metadata["extraction_method"] = "fallback_pypdf2"
            
            return content, metadata, []
            
        except Exception as e:
            logger.error(f"Fallback extraction failed: {e}")
            return "", metadata, []
    
    # Reuse existing extraction methods from original processor
    def _extract_contract_name(self, content: str, filename: str) -> Optional[str]:
        """Extract contract name from content or filename"""
        name_patterns = [
            r"Contract Title:?\\s*(.+)",
            r"Agreement:\\s*(.+)",
            r"CONTRACT:\\s*(.+)",
            r"^(.+(?:Agreement|Contract))",
        ]
        
        for pattern in name_patterns:
            match = re.search(pattern, content, re.IGNORECASE | re.MULTILINE)
            if match:
                name = match.group(1).strip()
                if len(name) < 100:
                    return name
        
        base_name = os.path.splitext(filename)[0]
        return base_name.replace('_', ' ').title()
    
    def _extract_parties(self, content: str) -> List[str]:
        """Extract contract parties from content"""
        parties = []
        
        party_patterns = [
            r"Party 1:?\\s*(.+?)(?:\\n|Party 2)",
            r"Party 2:?\\s*(.+?)(?:\\n|$)",
            r"Parties:?\\s*(.+?)(?:\\n\\n|\\nCONTRACT)",
            r"Between:?\\s*(.+?)\\s*and\\s*(.+?)(?:\\n|$)",
        ]
        
        for pattern in party_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE | re.DOTALL)
            for match in matches:
                if isinstance(match, tuple):
                    parties.extend(match)
                else:
                    parties.append(match)
        
        company_pattern = r'([A-Za-z\\s]+(?:Inc\\.|LLC|Corp\\.?|Corporation|Ltd\\.?|Limited))'
        companies = re.findall(company_pattern, content)
        parties.extend(companies)
        
        cleaned_parties = []
        for party in parties:
            cleaned = party.strip().replace('\\n', ' ')
            if cleaned and len(cleaned) > 3 and cleaned not in cleaned_parties:
                cleaned_parties.append(cleaned)
        
        return cleaned_parties[:10]
    
    def _extract_date(self, content: str, date_type: str) -> Optional[str]:
        """Extract dates from content"""
        date_patterns = {
            "start": [
                r"Start Date:?\\s*(\\d{4}-\\d{2}-\\d{2})",
                r"Effective Date:?\\s*(\\d{4}-\\d{2}-\\d{2})",
                r"Begin(?:ning)? Date:?\\s*(\\d{4}-\\d{2}-\\d{2})",
            ],
            "end": [
                r"End Date:?\\s*(\\d{4}-\\d{2}-\\d{2})",
                r"Expir(?:ation|y) Date:?\\s*(\\d{4}-\\d{2}-\\d{2})",
                r"Termination Date:?\\s*(\\d{4}-\\d{2}-\\d{2})",
            ],
            "renewal": [
                r"Renewal Date:?\\s*(\\d{4}-\\d{2}-\\d{2})",
                r"Review Date:?\\s*(\\d{4}-\\d{2}-\\d{2})",
            ]
        }
        
        patterns = date_patterns.get(date_type, [])
        for pattern in patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                return match.group(1)
        
        return None
    
    def _extract_key_clauses(self, content: str) -> List[str]:
        """Extract key clauses from content"""
        clauses = []
        
        clause_pattern = r"(\\d+\\.\\s*.+?)(?=\\n\\d+\\.|\\n[A-Z]{2,}|$)"
        matches = re.findall(clause_pattern, content, re.DOTALL)
        
        for match in matches:
            clause = match.strip()
            if len(clause) > 20 and len(clause) < 500:
                clauses.append(clause)
        
        return clauses[:10]
    
    def _extract_contact_info(self, content: str) -> Dict[str, Any]:
        """Extract contact information from content"""
        contact_info = {}
        
        email_pattern = r'\\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\\.[A-Z|a-z]{2,}\\b'
        emails = re.findall(email_pattern, content)
        if emails:
            contact_info["emails"] = list(set(emails))
        
        phone_pattern = r'\\+?1?[-.\\s]?\\(?[0-9]{3}\\)?[-.\\s]?[0-9]{3}[-.\\s]?[0-9]{4}'
        phones = re.findall(phone_pattern, content)
        if phones:
            contact_info["phones"] = list(set(phones))
        
        return contact_info
    
    def _extract_department(self, content: str, filename: str) -> Optional[str]:
        """Extract department information"""
        departments = ["Legal", "IT", "Finance", "Procurement", "Operations", "HR"]
        
        for dept in departments:
            if dept.lower() in content.lower() or dept.lower() in filename.lower():
                return dept
        
        return None
    
    def _extract_from_txt(self, file_path: str, metadata: Dict) -> tuple:
        """Extract content from TXT files (reuse from original)"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            metadata["extraction_method"] = "text_file"
            metadata["character_count"] = len(content)
            metadata["line_count"] = content.count('\\n')
            
            return content, metadata
            
        except UnicodeDecodeError:
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                    
                    metadata["extraction_method"] = f"text_file_{encoding}"
                    return content, metadata
                except UnicodeDecodeError:
                    continue
            
            logger.error(f"Failed to decode text file {file_path}")
            return "", metadata
    
    # Synchronous wrapper for backward compatibility
    def process_single_document(self, file_path: str, filename: str = None, 
                              extract_contracts: bool = True, 
                              custom_chunk_size: int = None,
                              metadata: Dict[str, Any] = None,
                              use_vision: bool = True) -> Dict[str, Any]:
        """Synchronous wrapper for document processing"""
        try:
            loop = asyncio.get_event_loop()
        except RuntimeError:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
        
        return loop.run_until_complete(
            self.process_document_async(
                file_path, filename, extract_contracts, 
                custom_chunk_size, metadata, use_vision
            )
        )